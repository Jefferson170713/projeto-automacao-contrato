import sys
import os
import pandas as pd
import locale
from num2words import num2words
from datetime import datetime
from docx import Document
from docx2pdf import convert

from PyQt5.QtGui import QIcon
locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')



from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QPushButton, QTableWidget, QTableWidgetItem,
    QFileDialog, QProgressBar, QLabel, QHBoxLayout, QGroupBox, QCheckBox, QButtonGroup
)
import pandas as pd

def resource_path(relative_path):
    import sys, os
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.df = pd.DataFrame()
        self.path_contract_aditivo = r'./Arquivos/CONTRATO ADITIVO PADRÃO.docx'
        self.path_contract_recisao_com_aviso = r'./Arquivos/DISTRATO MODELO ATUALIZADO - COM AVISO.docx'
        self.path_contract_recisao_sem_aviso = r'./Arquivos/DISTRATO MODELO ATUALIZADO - SEM AVISO.docx'
        self.path_csv = str()
        self.path_folder = str()
        self.list_remove_docx = []
        self.setWindowTitle('Aditivos')
        self.setWindowIcon(QIcon(r'./Arquivos/file-cloud.svg'))
        #self.setWindowIcon(QIcon(resource_path('Arquivos/hapvida_inside_circle.svg')))
        self.setGeometry(100, 100, 800, 600)
        
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.layout = QVBoxLayout()
        self.central_widget.setLayout(self.layout)

        # Grupo 1: barra de progresso em cima, 3 checkboxes embaixo
        self.group_checks = QGroupBox('Opções e Progresso')
        self.group_checks_layout = QVBoxLayout()
        self.group_checks.setLayout(self.group_checks_layout)
        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        self.group_checks_layout.addWidget(self.progress_bar)
        self.checkboxes_layout = QHBoxLayout()
        self.checkbox1 = QCheckBox('ADITIVO')
        self.checkbox2 = QCheckBox('RECISÃO COM AVISO')
        self.checkbox3 = QCheckBox('RECISÃO SEM AVISO')
        self.checkboxes_layout.addWidget(self.checkbox1)
        self.checkboxes_layout.addWidget(self.checkbox2)
        self.checkboxes_layout.addWidget(self.checkbox3)
        self.group_checks_layout.addLayout(self.checkboxes_layout)
        self.layout.addWidget(self.group_checks)

        # Grupo 2: só a tabela
        self.group_table = QGroupBox('📊 Contratos de Aditivos')
        self.group_table_layout = QVBoxLayout()
        self.group_table.setLayout(self.group_table_layout)
        self.table = QTableWidget()
        self.group_table_layout.addWidget(self.table)
        self.layout.addWidget(self.group_table)

        # Grupo 3: botões (já estava pronto)
        self.group_buttons = QGroupBox('🔍 Procurar')
        self.button_layout = QHBoxLayout()
        self.group_buttons.setLayout(self.button_layout)
        self.btn_select_file = QPushButton('📄 Selecionar arquivo CSV')
        self.btn_select_folder = QPushButton('📁 Selecionar pasta de saída')
        self.button_layout.addWidget(self.btn_select_file)
        self.button_layout.addWidget(self.btn_select_folder)
        self.layout.addWidget(self.group_buttons)

        self.btn_select_file.clicked.connect(self.abrir_e_carregar_csv)
        self.btn_select_folder.clicked.connect(self.selecionar_pasta_destino)

        self.checkbox1.setChecked(True)
        self.button_group = QButtonGroup(self)
        self.button_group.setExclusive(True)
        self.button_group.addButton(self.checkbox1)
        self.button_group.addButton(self.checkbox2)
        self.button_group.addButton(self.checkbox3)



    def ler_csv_personalizado(self, caminho_arquivo):
        """
        Lê um arquivo CSV com separador ';' e encoding 'latin1'.
        Atualiza a tabela com os dados lidos.
        """
        
        try:
            self.df = pd.read_csv(caminho_arquivo, sep=';', encoding='latin1')
        except Exception as erro:
            from PyQt5.QtWidgets import QMessageBox
            QMessageBox.critical(self, 'Erro ao ler CSV', str(erro))
            return
        # Atualiza a tabela
        self.atualiza_tabela()

    # função para atualizar a tabela
    def atualiza_tabela(self):
        self.table.setColumnCount(0)
        if not self.df.empty:
            self.table.setColumnCount(len(self.df.columns))
            self.table.setHorizontalHeaderLabels(self.df.columns.astype(str).tolist())
            self.table.setRowCount(len(self.df))
            for i, row in self.df.iterrows():
                for j, value in enumerate(row):
                    self.table.setItem(i, j, QTableWidgetItem(str(value)))

    def abrir_e_carregar_csv(self):
        options = QFileDialog.Options()
        options |= QFileDialog.ReadOnly
        file_path, _ = QFileDialog.getOpenFileName(self, 'Selecione o arquivo CSV', '', 'CSV Files (*.csv);;Todos Arquivos (*)', options=options)
        if file_path:
            self.ler_csv_personalizado(file_path)

    # função do valor descritivo
    def value_description(self, serie):
        valores_float = [float(v.replace("R$ ", "").replace(".", "").replace(",", ".")) for v in serie]
        serie = [num2words(v, lang='pt_BR', to='currency') for v in valores_float]
        return serie

    # função do valor do título
    def value_title(self, serie):
        return serie.str.capitalize()

    # função do valor monetário
    def value_money(self, valor):
        locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')
        valor_formatado = locale.format_string('%.2f', valor, grouping=True)
        return valor_formatado

    # função da data por extenso
    def data_por_extenso(self, data_str):
        # Converte a string em um objeto datetime
        data = datetime.strptime(data_str, "%d/%m/%Y")
        
        # Lista com os nomes dos meses
        meses = [
            "janeiro", "fevereiro", "março", "abril", "maio", "junho",
            "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"
        ]
        
        # Monta a data por extenso
        return f"{data.day} de {meses[data.month - 1]} de {data.year}"
    
    # função de reaguste do self.df
    def reajuste_df(self):
        self.df['NUM_ADITIVO'] = self.df['NUM_ADITIVO'].astype(str)

        self.df['VALOR_DESCRICAO'] = self.value_description(self.df['NOVO_VALOR'])
        self.df['VALOR_DESCRICAO'] = self.value_title(self.df['VALOR_DESCRICAO'])

        # Converta NOVO_VALOR para float antes de aplicar value_money
        self.df['NOVO_VALOR'] = self.df['NOVO_VALOR'].apply(lambda x: float(str(x).replace("R$ ", "").replace(".", "").replace(",", ".")))
        self.df['NOVO_VALOR'] = self.df['NOVO_VALOR'].apply(self.value_money)
        self.df['NOVO_VALOR'] = 'R$ ' + self.df['NOVO_VALOR']

        # datas agora
        self.df['DATA_CONTRATO_EXTENSO'] = self.df['DATA_CONTRATO'].apply(self.data_por_extenso)
        self.df['DATA_RETROAGE_EXTENSO'] = self.df['DATA_RETROAGE'].apply(self.data_por_extenso)

    # função para selecionar a pasta de destino
    def selecionar_pasta_destino(self):
        options = QFileDialog.Options()
        options |= QFileDialog.ReadOnly
        folder_path = QFileDialog.getExistingDirectory(self, 'Selecione a pasta de destino', options=options)
        if folder_path:
            self.path_folder = folder_path

        if self.checkbox1.isChecked():
            self.reajuste_df()
            self.criar_aditivos()

        if self.checkbox2.isChecked():
            self.criar_recisao_com_aviso()
            print("Temporary files removed:")
            self.list_remove_docx = self.remove_temp_files()

        if self.checkbox3.isChecked():
            self.criar_recisao_sem_aviso()
            print("Temporary files removed:")
            self.list_remove_docx = self.remove_temp_files()

    # função de criação dos aditivos
    def criar_aditivos(self):
        name_mes = self.df['MES'].iloc[0]
        bar_progresso = 0
        self.progress_bar.setValue(bar_progresso)
        bar_progresso_max = len(self.df)

        for num, nome in enumerate(self.df['NOME']):
            dict_modification = {
                "@num_aditivo@" : self.df['NUM_ADITIVO'].iloc[num],
                "@razao_social@" : self.df['RAZAO_SOCIAL'].iloc[num],
                "@cnpj@" : self.df['CNPJ'].iloc[num],
                "@cidade@" : self.df['CIDADE'].iloc[num],
                "@endereco@" : self.df['ENDERECO'].iloc[num],
                "@numero@" : self.df['NUMERO'].iloc[num],
                "@bairro@" : self.df['BAIRRO'].iloc[num],
                "data_contrato@" : self.df['DATA_CONTRATO_EXTENSO'].iloc[num],
                "@data_retroagem@" : self.df['DATA_RETROAGE_EXTENSO'].iloc[num],
                "@valor@" : self.df['NOVO_VALOR'].iloc[num],
                "@valor_descrito@" : self.df['VALOR_DESCRICAO'].iloc[num],
            }

            doc = Document(self.path_contract_aditivo)
            for para in doc.paragraphs:
                for run in para.runs:
                    for key, value in dict_modification.items():
                        if key in run.text:
                            run.text = run.text.replace(key, str(value))
                            
            name_arquivo = f'CONTRATO ADITIVO - {nome} - {name_mes}.docx'
            caminho_completo = os.path.join(self.path_folder, name_arquivo)
            doc.save(caminho_completo)
            bar_progresso += 1
            self.progress_bar.setValue(int((bar_progresso / bar_progresso_max) * 100))

        self.progress_bar.setValue(100)

    def criar_recisao_com_aviso(self):
        bar_progresso = 0
        self.progress_bar.setValue(bar_progresso)
        bar_progresso_max = len(self.df)

        for num, razao_social in enumerate(self.df['RAZAO_SOCIAL']):
            dict_modification = {
                "XRAZAO_SOCIALX" : self.df['RAZAO_SOCIAL'].iloc[num],
                "XDATA_1X" : self.df['DATA_DESLIGAMENTO'].iloc[num],
                "XDATA_2X": self.df['DATA_CONTRATACAO'].iloc[num],
                "XCNPJX": self.df['CNPJ'].iloc[num],
                "XAVISOX": self.df['DIAS_EFEITO'].iloc[num],
                "XCIDADEX": self.df['CIDADE'].iloc[num],
                "XENDERECOX": self.df['ENDERECO'].iloc[num],
                "XLOCAL_NUMEROX": self.df['NUMERO'].iloc[num],
                "XBAIRROX": self.df['BAIRRO'].iloc[num],
                "XCEPX": self.df['CEP'].iloc[num],
            }

            nome_file_save = self.df['NOME'].iloc[num]

            doc = Document(self.path_contract_recisao_com_aviso)

            for para in doc.paragraphs:
                for run in para.runs:
                    for key, value in dict_modification.items():
                        if key in run.text:
                            run.text = run.text.replace(key, str(value))

            name_arquivo_docx = f'RECISÃO CONTRATUAL - HAP - NOTIFICAÇÃO COM AVISO - {nome_file_save}.docx'
            name_arquivo_pdf = f'RECISÃO CONTRATUAL - HAP - NOTIFICAÇÃO COM AVISO - {nome_file_save}.pdf'
            caminho_completo_docx = os.path.join(self.path_folder, name_arquivo_docx)
            caminho_completo_pdf = os.path.join(self.path_folder, name_arquivo_pdf)
            doc.save(caminho_completo_docx)
            self.list_remove_docx.append(caminho_completo_docx)
            bar_progresso += 1
            self.progress_bar.setValue(int((bar_progresso / bar_progresso_max) * 100))

        self.progress_bar.setValue(100)


    def criar_recisao_sem_aviso(self):
        bar_progresso = 0
        self.progress_bar.setValue(bar_progresso)
        bar_progresso_max = len(self.df)

        for num, razao_social in enumerate(self.df['RAZAO_SOCIAL']):
            dict_modification = {
                "XRAZAO_SOCIALX" : self.df['RAZAO_SOCIAL'].iloc[num],
                "XDATA_1X" : self.df['DATA_DESLIGAMENTO'].iloc[num],
                "XDATA_2X": self.df['DATA_CONTRATACAO'].iloc[num],
                "XCNPJX": self.df['CNPJ'].iloc[num],
                # "XAVISOX": self.df['DIAS_EFEITO'].iloc[num],
                "XCIDADEX": self.df['CIDADE'].iloc[num],
                "XENDERECOX": self.df['ENDERECO'].iloc[num],
                "XLOCAL_NUMEROX": self.df['NUMERO'].iloc[num],
                "XBAIRROX": self.df['BAIRRO'].iloc[num],
                "XCEPX": self.df['CEP'].iloc[num],
            }

            name_file_save = self.df['NOME'].iloc[num]

            doc = Document(self.path_contract_recisao_sem_aviso)

            for para in doc.paragraphs:
                for run in para.runs:
                    for key, value in dict_modification.items():
                        if key in run.text:
                            run.text = run.text.replace(key, str(value))

            name_arquivo_docx = f'RECISÃO CONTRATUAL - HAP - NOTIFICAÇÃO SEM AVISO - {name_file_save}.docx'
            name_arquivo_pdf = f'RECISÃO CONTRATUAL - HAP - NOTIFICAÇÃO SEM AVISO - {name_file_save}.pdf'
            caminho_completo_docx = os.path.join(self.path_folder, name_arquivo_docx)
            # caminho_completo_pdf = os.path.join(self.path_folder, name_arquivo_pdf)
            doc.save(caminho_completo_docx)
            self.list_remove_docx.append(caminho_completo_docx)
            bar_progresso += 1
            self.progress_bar.setValue(int((bar_progresso / bar_progresso_max) * 100))

        self.progress_bar.setValue(100)

    def remove_temp_files(self):
        for num, caminho in enumerate(self.list_remove_docx):
            print(f'{num} - {caminho}')
        list_vazia = []
        return list_vazia

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())
